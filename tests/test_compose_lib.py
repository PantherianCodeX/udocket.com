from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, Mapping, Tuple
import zipfile

from docx import Document as DocxDocument

import pytest

from packages.udocket_core.agents.compose_lib import (
    ComposeAgent,
    ComposeConfig,
    ComposeResult,
    ComposeStageError,
    ComposeState,
    GuardReport,
    LaneActionDirective,
    LaneOutcome,
    LaneRuntimeState,
    _factuality_report,
    _stable_doc_fingerprint,
    ComposeInputs,
    LANE_CONFIGS,
)
from packages.udocket_core.agents.compose.state import QAReviewerResult, _merge_lane_outcomes
from packages.udocket_core.agents.compose.orchestrator import ComposeOrchestrator
from packages.udocket_core.json_utils import JSONObject
from packages.udocket_core.agents.compose.errors import ComposeStageError
from tests._typing import MonkeyPatch


ClientResponse = Tuple[str, Dict[str, int], str, str]


def _make_fake_qa_step(
    staff_report_text: str,
) -> Callable[[ComposeOrchestrator, ComposeState, Mapping[str, JSONObject], Optional[Callable[[str, str, JSONObject], None]]], dict[str, object]]:
    def _fake_qa_step(
        self: ComposeOrchestrator,
        *,
        state: ComposeState,
        provider_credentials: Mapping[str, JSONObject],
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> dict[str, object]:
        if "client" not in state.lanes:
            try:
                state.lanes["client"] = state.client.to_outcome()
            except ComposeStageError:
                pass
        if "lawyer" not in state.lanes:
            try:
                state.lanes["lawyer"] = state.lawyer.to_outcome()
            except ComposeStageError:
                pass
        qa_result = QAReviewerResult(
            status="ok",
            alerts=[],
            recommendations=[],
            staff_report=staff_report_text,
            provider="stub",
            lane_actions={
                "client": LaneActionDirective(action="none"),
                "lawyer": LaneActionDirective(action="none"),
            },
            global_notes="",
        )
        state.qa = qa_result
        state.qa_iterations += 1
        return {
            "qa": qa_result,
            "qa_iterations": state.qa_iterations,
            "stage_usage": {"compose.qa_reviewer": {"prompt_tokens": 0, "completion_tokens": 0}},
        }

    return _fake_qa_step


CLIENT_VALID_DOC = """## Case Overview
At [00:01] the judge confirmed the interim custody order stayed active for both parents. Counsel at [00:01] reminded everyone that earlier schedules still govern each exchange. At [02:00] the clerk reviewed disclosure timelines and reiterated expectations for continued compliance. At [02:00] the judge requested frequent progress updates so the next conference remains productive. Counsel also noted at [02:00] that maintaining predictable routines helps the children adapt calmly.

## Key People and Roles
Alex at [00:01] described handling school drop-offs, medical visits, and bedtime routines every week. He added at [00:01] that grandparents assist daily so obligations remain balanced. Morgan at [02:00] confirmed disclosure tasks are pending but promised detailed weekly reports. She stated at [02:00] that extended family coordinate travel to avoid missed exchanges. Counsel at [02:00] committed to logging every support person involved with the schedule.

## Timeline of Events
At [00:01] the hearing opened with a recap of earlier interim orders and compliance history. Around [00:30] counsel listed affidavits and exhibits that still awaited review. At [01:10] the parties described cooperative steps and flagged remaining disclosure gaps from the prior month. Near [01:40] the judge summarised expectations for document exchanges before upcoming deadlines. By [02:00] the court restated the operative order and mapped deadlines for the next appearance.

## Main Issues
At [00:01] the court questioned whether interim custody terms should change before the next conference. At [02:00] compliance delays raised concerns about readiness for motions and long term settlement tasks. Counsel also noted at [02:00] that any adjustment must respect the children’s routine stability.

## Next Steps / Preparation Notes
At [02:00] counsel must gather financial statements, education reports, and parenting calendars before the next conference. Teams should log cooperation noted at [00:01] to keep status updates accurate. Everyone prepares concise notes linking each outstanding obligation to transcript timestamps such as [02:00] before filing. The parties also agreed at [02:00] to share weekly progress summaries with the clerk.
"""

LAWYER_VALID_DOC = """## Case Summary
This matter involves interim custody arrangements reviewed on January 5, 2024 at [00:01], where the court confirmed the existing order remained in force pending additional disclosure. The transcript at [02:00] records the judge’s expectation that both parties continue cooperation while preparing for the next settlement-focused appearance, emphasising accuracy and timeliness.

## Parties and Roles
Alex, the Applicant, presented evidence about parenting schedules at [00:01] and outlined travel logistics for school and extracurricular commitments. Morgan, the Respondent, discussed at [02:00] how disclosure will be completed, identified supporting relatives who assist with exchanges, and confirmed continued participation in case management discussions.

## Factual Background
The hearing at [00:01] opened with the judge recounting prior orders, including the interim custody decision that balanced weekday and weekend responsibilities between the parties. Counsel reviewed affidavits, financial statements, and parenting time notes at [00:45] to confirm the written record matched the oral conditions previously delivered. Shortly afterward, around [00:45], the parties debated outstanding document production, prompting the court to reiterate at [00:45] that disclosure must remain on schedule. By [01:20], Alex described daily routines, medical appointments, and childcare arrangements that depended on the interim terms, while Morgan acknowledged at [01:20] the need to supplement evidence with updated income details. At [02:00] the judge summarised the expectations: maintain the interim order, exchange missing disclosure promptly, and prepare a comprehensive joint update before the next conference. The transcript at [02:00] further notes that both parties accepted these directives without objection, providing a stable framework for continued negotiations.

## Issues Presented
Key issues include whether any material change since [00:01] justifies altering the interim custody order, and how the disclosure delays highlighted at [02:00] affect readiness for conferences or motions. The court also needs clarity at [02:00] on scheduling conflicts and the feasibility of future attendance dates given the family’s logistical commitments.

## Evidence / Supporting Facts
Transcript excerpts at [00:01], [00:45], and [02:00] document the judge’s oral reasons and the parties’ acknowledgements of ongoing duties. Counsel referenced undertakings, email correspondence, and previously filed statements at [02:00] to substantiate compliance, and they agreed to provide supplemental summaries tying each obligation to specific exhibits and time-stamped transcript passages.

## Procedural Status / Next Known Steps
The case conference remains scheduled following the directives issued at [02:00], and the parties must exchange updated financial materials, medical summaries, and childcare plans beforehand. The court also requested written updates at [02:00] detailing progress on disclosure, coordination with support services, and confirmation that interim arrangements continue without disruption.
"""


def _write_inputs(base_dir: Path) -> tuple[Path, Path, Path]:
    summary_json = base_dir / "summary.json"
    summary_json.write_text(
        json.dumps(
            {
                "parties": [
                    {"name": "Alex", "role": "Applicant"},
                    {"name": "Morgan", "role": "Respondent"},
                ],
                "facts": [
                    {"text": "Hearing held on January 5, 2024."},
                    {"text": "Interim custody order currently in effect."},
                ],
            }
        ),
        encoding="utf-8",
    )

    summary_md = base_dir / "summary.md"
    summary_md.write_text(
        """# Summary

Alex and Morgan appeared before the court on January 5, 2024.
The court issued an interim order.
""",
        encoding="utf-8",
    )

    timeline_json = base_dir / "timeline.json"
    timeline_json.write_text(
        json.dumps(
            {
                "events": [
                    {
                        "id": "event-1",
                        "summary": "Hearing commenced",
                        "ts_start": 1.0,
                        "ts_end": 2.0,
                        "speakers": ["Alex"],
                        "references": ["[00:01]"]
                    },
                    {
                        "id": "event-2",
                        "summary": "Order granted",
                        "ts_start": 120.0,
                        "ts_end": 121.0,
                        "speakers": ["Judge"],
                        "references": ["[02:00]"]
                    },
                ]
            }
        ),
        encoding="utf-8",
    )

    return summary_json, summary_md, timeline_json


def test_compose_agent_parallel_lanes(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    case_dir = tmp_path / "case"
    docs_dir = case_dir / "docs"
    ops_dir = case_dir / "ops"
    docs_dir.mkdir(parents=True)
    ops_dir.mkdir(parents=True)

    summary_json, summary_md, timeline_json = _write_inputs(docs_dir)

    config = ComposeConfig(
        provider_chain=["stub"],
        temperature=0.2,
        lawyer_temperature=0.2,
        max_output_tokens=2048,
        max_client_attempts=2,
        max_lawyer_attempts=2,
        min_timestamp_references=1,
        qa_required=True,
        debug=True,
    )
    agent = ComposeAgent(config)

    client_doc = CLIENT_VALID_DOC

    lawyer_doc = LAWYER_VALID_DOC

    def fake_invoke(
        *,
        stage: str,
        system_prompt: str,
        user_prompt: str,
        temperature: float,
        provider_credentials: Mapping[str, JSONObject],
        config: ComposeConfig,
        settings: object,
    ) -> ClientResponse:
        if stage == "compose.client.draft":
            return client_doc, {"prompt_tokens": 100, "completion_tokens": 200}, "stub", "stub-model"
        if stage == "compose.client.revise":
            revised_client = client_doc.replace(
                "## Next Steps / Preparation Notes",
                "## Next Steps / Preparation Notes\nUpdated after QA.",
            )
            return revised_client, {"prompt_tokens": 90, "completion_tokens": 190}, "stub", "stub-model"
        if stage == "compose.lawyer.draft":
            return lawyer_doc, {"prompt_tokens": 120, "completion_tokens": 210}, "stub", "stub-model"
        if stage == "compose.lawyer.revise":
            revised_lawyer = lawyer_doc.replace(
                "## Procedural Status / Next Known Steps",
                "## Procedural Status / Next Known Steps\nQA confirmed clarity.",
            )
            return revised_lawyer, {"prompt_tokens": 110, "completion_tokens": 205}, "stub", "stub-model"
        if stage == "compose.qa_reviewer":
            response = json.dumps(
                {
                    "status": "ok",
                    "alerts": [],
                    "recommendations": [],
                    "staff_report": "# Staff Report\n\nAll checks passed.",
                    "global_notes": "",
                    "lane_actions": {
                        "client": {"action": "none", "revision_brief": ""},
                        "lawyer": {"action": "none", "revision_brief": ""},
                    },
                }
            )
            return response, {"prompt_tokens": 80, "completion_tokens": 40}, "stub", "stub-model"
        raise AssertionError(f"Unexpected stage: {stage}")

    for target in (
        "packages.udocket_core.agents.compose.orchestrator.invoke_llm",
        "packages.udocket_core.agents.compose.llm_runtime.invoke_llm",
        "packages.udocket_core.agents.compose.qa.invoke_llm",
    ):
        monkeypatch.setattr(target, fake_invoke)

    fake_qa_step = _make_fake_qa_step("# Staff Report\n\nAll checks passed.")
    monkeypatch.setattr(ComposeOrchestrator, "_qa_reviewer_step", fake_qa_step)

    result: ComposeResult = agent.compose(
        case_id="CASE-001",
        case_dir=case_dir,
        job_id="JOB-001",
        summary_json_path=summary_json,
        summary_markdown_path=summary_md,
        timeline_seed_path=timeline_json,
        entity_hint_path=None,
    )

    assert result.status == "ok"
    artifacts = result.artifacts
    bundle_path = artifacts.bundle_path
    assert bundle_path is not None
    assert bundle_path.exists()
    bundle_text = bundle_path.read_text(encoding="utf-8")
    assert "Part 1 – Client Summary" in bundle_text
    assert "Part 2 – Lawyer Brief" in bundle_text
    assert artifacts.client_markdown and artifacts.client_markdown.exists()
    assert artifacts.lawyer_markdown and artifacts.lawyer_markdown.exists()
    assert artifacts.staff_report and artifacts.staff_report.exists()
    assert artifacts.qa_report and artifacts.qa_report.exists()
    qa_text = artifacts.qa_report.read_text(encoding="utf-8")
    assert "## Client Lane" in qa_text
    assert result.meta_json.exists()
    meta = json.loads(result.meta_json.read_text(encoding="utf-8"))
    assert meta["client_attempts"] == 1
    assert meta["lawyer_attempts"] == 1
    assert meta["stage_usage"]
    assert meta["staff_report"] and Path(meta["staff_report"]).exists()


def test_compose_agent_revision_cycle(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    case_dir = tmp_path / "case"
    docs_dir = case_dir / "docs"
    ops_dir = case_dir / "ops"
    docs_dir.mkdir(parents=True)
    ops_dir.mkdir(parents=True)

    summary_json, summary_md, timeline_json = _write_inputs(docs_dir)

    config = ComposeConfig(
        provider_chain=["stub"],
        temperature=0.2,
        lawyer_temperature=0.2,
        max_output_tokens=2048,
        max_client_attempts=2,
        max_lawyer_attempts=1,
        min_timestamp_references=1,
        qa_required=True,
        debug=True,
    )
    agent = ComposeAgent(config)

    call_counts: dict[str, int] = {"compose.client.draft": 0, "compose.client.revise": 0}

    bad_client_doc = """## Case Overview
This section omits required references and headings.

## Key People and Roles
- Alex: Applicant

## Main Issues
- Missing sections cause failure.
"""

    good_client_doc = CLIENT_VALID_DOC

    lawyer_doc = LAWYER_VALID_DOC

    def fake_invoke(
        *,
        stage: str,
        system_prompt: str,
        user_prompt: str,
        temperature: float,
        provider_credentials: Mapping[str, JSONObject],
        config: ComposeConfig,
        settings: object,
    ) -> ClientResponse:
        if stage == "compose.client.draft":
            call_counts[stage] += 1
            return bad_client_doc, {"prompt_tokens": 50, "completion_tokens": 50}, "stub", "stub-model"
        if stage == "compose.client.revise":
            call_counts[stage] += 1
            return good_client_doc, {"prompt_tokens": 60, "completion_tokens": 70}, "stub", "stub-model"
        if stage == "compose.lawyer.draft":
            return lawyer_doc, {"prompt_tokens": 40, "completion_tokens": 60}, "stub", "stub-model"
        if stage == "compose.lawyer.revise":
            return lawyer_doc, {"prompt_tokens": 45, "completion_tokens": 65}, "stub", "stub-model"
        if stage == "compose.qa_reviewer":
            response = json.dumps(
                {
                    "status": "ok",
                    "alerts": [],
                    "recommendations": [],
                    "staff_report": "# Staff Report\n\nClient lane fixed.",
                    "global_notes": "",
                    "lane_actions": {
                        "client": {"action": "none", "revision_brief": ""},
                        "lawyer": {"action": "none", "revision_brief": ""},
                    },
                }
            )
            return response, {"prompt_tokens": 30, "completion_tokens": 30}, "stub", "stub-model"
        raise AssertionError(f"Unexpected stage: {stage}")

    for target in (
        "packages.udocket_core.agents.compose.orchestrator.invoke_llm",
        "packages.udocket_core.agents.compose.llm_runtime.invoke_llm",
        "packages.udocket_core.agents.compose.qa.invoke_llm",
    ):
        monkeypatch.setattr(target, fake_invoke)

    fake_qa_step = _make_fake_qa_step("# Staff Report\n\nClient lane fixed.")
    monkeypatch.setattr(ComposeOrchestrator, "_qa_reviewer_step", fake_qa_step)
    assert ComposeOrchestrator._qa_reviewer_step is fake_qa_step

    result = agent.compose(
        case_id="CASE-REV",
        case_dir=case_dir,
        job_id="JOB-REV",
        summary_json_path=summary_json,
        summary_markdown_path=summary_md,
        timeline_seed_path=timeline_json,
        entity_hint_path=None,
    )

    meta = json.loads(result.meta_json.read_text(encoding="utf-8"))
    assert meta["client_attempts"] == 2
    assert meta["lawyer_attempts"] == 1
    assert call_counts["compose.client.draft"] == 1
    assert call_counts["compose.client.revise"] == 1
    bundle_path = result.artifacts.bundle_path
    assert bundle_path is not None
    bundle_text = bundle_path.read_text(encoding="utf-8")
    assert "Next Steps / Preparation Notes" in bundle_text


def test_compose_agent_docx_template(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    case_dir = tmp_path / "case"
    docs_dir = case_dir / "docs"
    ops_dir = case_dir / "ops"
    docs_dir.mkdir(parents=True)
    ops_dir.mkdir(parents=True)

    summary_json, summary_md, timeline_json = _write_inputs(docs_dir)

    template_path = tmp_path / "compose_template.docx"
    template_doc = DocxDocument()
    template_doc.add_heading("Client Summary", level=2)
    template_doc.add_paragraph("{{ client_summary }}")
    template_doc.add_heading("Lawyer Brief", level=2)
    template_doc.add_paragraph("{{ lawyer_brief }}")
    template_doc.add_heading("Staff Notes", level=2)
    template_doc.add_paragraph("{{ staff_report_plain }}")
    template_doc.save(str(template_path))

    config = ComposeConfig(
        provider_chain=["stub"],
        temperature=0.2,
        lawyer_temperature=0.2,
        max_output_tokens=2048,
        max_client_attempts=2,
        max_lawyer_attempts=2,
        min_timestamp_references=1,
        qa_required=True,
        debug=True,
        doc_template_path=template_path,
    )
    agent = ComposeAgent(config)

    client_doc = CLIENT_VALID_DOC

    lawyer_doc = LAWYER_VALID_DOC

    def fake_invoke(
        *,
        stage: str,
        system_prompt: str,
        user_prompt: str,
        temperature: float,
        provider_credentials: Mapping[str, JSONObject],
        config: ComposeConfig,
        settings: object,
    ) -> ClientResponse:
        if stage == "compose.client.draft":
            return client_doc, {"prompt_tokens": 80, "completion_tokens": 120}, "stub", "stub-model"
        if stage == "compose.client.revise":
            revised_client = client_doc + "\nRevision note added."
            return revised_client, {"prompt_tokens": 70, "completion_tokens": 110}, "stub", "stub-model"
        if stage == "compose.lawyer.draft":
            return lawyer_doc, {"prompt_tokens": 90, "completion_tokens": 130}, "stub", "stub-model"
        if stage == "compose.lawyer.revise":
            revised_lawyer = lawyer_doc + "\nRevision note recorded."
            return revised_lawyer, {"prompt_tokens": 85, "completion_tokens": 120}, "stub", "stub-model"
        if stage == "compose.qa_reviewer":
            response = json.dumps(
                {
                    "status": "ok",
                    "alerts": [],
                    "recommendations": [],
                    "staff_report": "# Staff Report\n\nAll checks passed.",
                    "global_notes": "",
                    "lane_actions": {
                        "client": {"action": "none", "revision_brief": ""},
                        "lawyer": {"action": "none", "revision_brief": ""},
                    },
                }
            )
            return response, {"prompt_tokens": 60, "completion_tokens": 30}, "stub", "stub-model"
        raise AssertionError(stage)

    for target in (
        "packages.udocket_core.agents.compose.orchestrator.invoke_llm",
        "packages.udocket_core.agents.compose.llm_runtime.invoke_llm",
        "packages.udocket_core.agents.compose.qa.invoke_llm",
    ):
        monkeypatch.setattr(target, fake_invoke)

    fake_qa_step = _make_fake_qa_step("# Staff Report\n\nAll checks passed.")
    monkeypatch.setattr(ComposeOrchestrator, "_qa_reviewer_step", fake_qa_step)

    result = agent.compose(
        case_id="CASE-003",
        case_dir=case_dir,
        job_id="JOB-003",
        summary_json_path=summary_json,
        summary_markdown_path=summary_md,
        timeline_seed_path=timeline_json,
        entity_hint_path=None,
    )

    client_docx_path = result.artifacts.client_docx
    assert client_docx_path is not None
    assert client_docx_path.exists()
    client_render = DocxDocument(str(client_docx_path))
    client_text = "\n".join(paragraph.text for paragraph in client_render.paragraphs)
    assert "Client Summary" in client_text
    staff_report_path = result.artifacts.staff_report
    assert staff_report_path is not None and staff_report_path.exists()
    staff_text = staff_report_path.read_text(encoding="utf-8")
    assert "All checks passed." in staff_text
    assert "{{" not in client_text
    with zipfile.ZipFile(client_docx_path, "r") as zf:
        client_xml = zf.read("word/document.xml").decode("utf-8")
    assert "the judge confirmed the interim custody order stayed active" in client_xml.lower()

    lawyer_docx_path = result.artifacts.lawyer_docx
    assert lawyer_docx_path is not None
    assert lawyer_docx_path.exists()
    lawyer_render = DocxDocument(str(lawyer_docx_path))
    lawyer_text = "\n".join(paragraph.text for paragraph in lawyer_render.paragraphs)
    assert "Lawyer Brief" in lawyer_text
    assert "All checks passed." in lawyer_text
    assert "{{" not in lawyer_text
    with zipfile.ZipFile(lawyer_docx_path, "r") as zf:
        lawyer_xml = zf.read("word/document.xml").decode("utf-8")
    assert "interim custody arrangements reviewed" in lawyer_xml


def _guard_ok() -> GuardReport:
    return GuardReport(ok=True, errors=[], warnings=[], checks={})


def test_merge_lane_outcomes_handles_removals_and_replacements() -> None:
    client_outcome = LaneOutcome(
        document="client",
        structure_report=_guard_ok(),
        compliance_report=_guard_ok(),
        factuality_report=_guard_ok(),
        attempts=1,
        history=[],
        stage_usage={"stage": {"tokens": 10}},
        token_usage={"tokens": 10},
        providers=["stub"],
        models=["model"],
        stage_durations={},
    )
    lawyer_outcome = LaneOutcome(
        document="lawyer",
        structure_report=_guard_ok(),
        compliance_report=_guard_ok(),
        factuality_report=_guard_ok(),
        attempts=1,
        history=[],
        stage_usage={"stage": {"tokens": 20}},
        token_usage={"tokens": 20},
        providers=["stub"],
        models=["model"],
        stage_durations={},
    )
    existing = {"client": client_outcome, "lawyer": lawyer_outcome}
    removed = _merge_lane_outcomes(existing, {"client": None})
    assert "client" not in removed
    assert removed["lawyer"] is lawyer_outcome

    replacement = LaneOutcome(
        document="client-new",
        structure_report=_guard_ok(),
        compliance_report=_guard_ok(),
        factuality_report=_guard_ok(),
        attempts=2,
        history=[],
        stage_usage={},
        token_usage={},
        providers=["stub"],
        models=["model"],
        stage_durations={},
    )
    replaced = _merge_lane_outcomes(existing, {"client": replacement})
    assert replaced["client"] is replacement
    assert replaced["lawyer"] is lawyer_outcome


def test_qa_iteration_limit_enforced() -> None:
    config = ComposeConfig(provider_chain=["stub"], qa_iteration_limit=1)
    agent = ComposeAgent(config)
    inputs = ComposeInputs(
        summary_markdown="",
        summary_data={},
        timeline_seeds=[],
        entity_hints={},
        intake={},
        case_metadata={},
    )
    client_state = LaneRuntimeState(
        lane="client", config=LANE_CONFIGS["client"], max_attempts=1
    )
    lawyer_state = LaneRuntimeState(
        lane="lawyer", config=LANE_CONFIGS["lawyer"], max_attempts=1
    )
    state = ComposeState(inputs=inputs, client=client_state, lawyer=lawyer_state)
    state.qa_iterations = 1

    with pytest.raises(ComposeStageError):
        agent._qa_reviewer_step(
            state=state,
            provider_credentials={},
            progress=None,
        )


def test_editor_rejects_unchanged_document(monkeypatch: MonkeyPatch) -> None:
    config = ComposeConfig(provider_chain=["stub"], enable_editor=True)
    agent = ComposeAgent(config)
    base_doc = "## Heading\n\nContent with [00:01] reference."
    inputs = ComposeInputs(
        summary_markdown="",
        summary_data={},
        timeline_seeds=[],
        entity_hints={},
        intake={},
        case_metadata={},
    )
    client_state = LaneRuntimeState(
        lane="client", config=LANE_CONFIGS["client"], max_attempts=2
    )
    client_state.document = base_doc
    client_state.last_document_hash = _stable_doc_fingerprint(base_doc)

    lawyer_state = LaneRuntimeState(
        lane="lawyer", config=LANE_CONFIGS["lawyer"], max_attempts=2
    )
    state = ComposeState(inputs=inputs, client=client_state, lawyer=lawyer_state)

    def fake_invoke(
        *,
        stage: str,
        system_prompt: str,
        user_prompt: str,
        temperature: float,
        provider_credentials: Mapping[str, JSONObject],
        config: ComposeConfig,
        settings: object,
    ) -> ClientResponse:
        response = json.dumps({"document": base_doc, "change_log": []})
        return response, {"prompt_tokens": 5, "completion_tokens": 5}, "stub", "stub-model"

    for target in (
        "packages.udocket_core.agents.compose.orchestrator.invoke_llm",
        "packages.udocket_core.agents.compose.llm_runtime.invoke_llm",
        "packages.udocket_core.agents.compose.qa.invoke_llm",
    ):
        monkeypatch.setattr(target, fake_invoke)

    monkeypatch.setattr(
        "packages.udocket_core.agents.compose.orchestrator.ComposeOrchestrator._qa_reviewer_step",
        fake_qa_step,
    )
    directive = LaneActionDirective(action="editor", revision_brief="Clarify tone.")

    agent._run_lane_editor(
        state=state,
        lane="client",
        directive=directive,
        provider_credentials={},
        progress=None,
    )
    assert state.client.document == base_doc
    assert state.client.editor_attempted is True
    assert state.client.providers[-1] == "stub"


def test_factuality_report_accepts_extended_timestamps_and_ids() -> None:
    document = """## Case Overview
Event event-1 discussed at [00:01] with context.
Further details about evt-2 appear at [01:02:03].
"""
    report = _factuality_report(
        document,
        claimable_atoms=[],
        timeline_events=[
            {"id": "event-1"},
            {"id": "evt-2"},
        ],
        min_timestamp_references=2,
    )
    assert report.ok
    assert not report.errors
    assert not report.warnings
