from __future__ import annotations

# pyright: strict

import json
import logging
import os
import re
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Callable, Iterable, Mapping, Optional, Sequence, cast

from docx import Document  # type: ignore[import]
from langgraph.graph import END, START, StateGraph  # type: ignore[import]

from packages.udocket_core.json_utils import (
    JSONObject,
    JSONValue,
    coerce_float,
    coerce_json_object,
    coerce_object_list,
    coerce_str,
    coerce_str_list,
    load_json_object,
    load_json_value,
    write_json_object,
)

from .common import append_jsonl, ensure_dir, next_versioned
from .common.docx import write_basic_docx
from ..llm import LLMSettings, load_llm_settings
from ..llm.runtime import ChatClientError, build_chat_client, build_provider_runtime_config
from .compose.llm_profiles import (
    DEFAULT_LAWYER_TEMPERATURE,
    DEFAULT_MAX_OUTPUT_TOKENS,
    DEFAULT_TEMPERATURE,
    LANE_CONFIGS,
    LaneConfig,
    QA_REVIEWER_SYSTEM_PROMPT,
    STAGE_MODEL_DEFAULTS,
    CLIENT_DRAFT_USER_INSTRUCTION,
    CLIENT_REVISION_USER_INSTRUCTION,
    LAWYER_DRAFT_USER_INSTRUCTION,
    LAWYER_REVISION_USER_INSTRUCTION,
    REVISION_HEADER_TEMPLATE,
    lane_system_prompt,
)


logger = logging.getLogger("udocket.compose.agent")


DEFAULT_PROVIDER_CHAIN: list[str] = ["azure"]

DOC_TEMPLATE_ENV = "COMPOSE_DOCX_TEMPLATE"


QA_REVIEWER_STATUS_OK = {"ok", "pass", "approved"}


class ComposeStageError(RuntimeError):
    def __init__(self, stage: str, message: str) -> None:
        super().__init__(f"{stage}: {message}")
        self.stage = stage


def _truthy(value: Optional[str], default: bool) -> bool:
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _safe_float(value: Optional[str], fallback: float) -> float:
    try:
        return float(value) if value else fallback
    except (TypeError, ValueError):
        return fallback


def _safe_int(value: Optional[str], fallback: int) -> int:
    try:
        return int(value) if value else fallback
    except (TypeError, ValueError):
        return fallback


def _normalize_providers(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for raw in values:
        name = (raw or "").strip().lower()
        if not name or name in seen:
            continue
        seen.add(name)
        ordered.append(name)
    return ordered


def _coerce_optional_object(value: JSONValue) -> JSONObject:
    return coerce_json_object(value) if isinstance(value, Mapping) else {}


def _collect_alias_items(source: Mapping[str, JSONValue], *keys: str) -> list[JSONObject]:
    items: list[JSONObject] = []
    for key in keys:
        payload = source.get(key)
        if isinstance(payload, Mapping):
            items.append(coerce_json_object(payload))
        elif isinstance(payload, Sequence):
            for element in payload:
                if isinstance(element, Mapping):
                    items.append(coerce_json_object(element))
    return items


def _extract_parties(summary_data: Mapping[str, JSONValue]) -> list[JSONObject]:
    parties_value = summary_data.get("parties")
    parties: list[JSONObject] = []
    if isinstance(parties_value, Mapping):
        for role_key in (
            "client",
            "applicant",
            "plaintiff",
            "petitioner",
            "respondent",
            "defendant",
            "opposing",
        ):
            role_val = parties_value.get(role_key)
            if isinstance(role_val, Mapping):
                obj = coerce_json_object(role_val)
                obj.setdefault("role", role_key.upper())
                parties.append(obj)
        for collection_key in ("counsel", "lawyers", "representatives"):
            for item in coerce_object_list(parties_value.get(collection_key)):
                obj = coerce_json_object(item)
                obj.setdefault("role", collection_key[:-1].upper())
                parties.append(obj)
    elif isinstance(parties_value, Sequence):
        for element in parties_value:
            if isinstance(element, Mapping):
                parties.append(coerce_json_object(element))
    if not parties:
        parties.extend(_collect_alias_items(summary_data, "parties"))
    return parties


def _extract_deadlines(summary_data: Mapping[str, JSONValue]) -> list[JSONObject]:
    return _collect_alias_items(summary_data, "deadlines", "upcoming_deadlines")


def _extract_orders(summary_data: Mapping[str, JSONValue]) -> list[JSONObject]:
    orders = _collect_alias_items(summary_data, "orders", "orders_and_directions")
    if not orders:
        orders = _collect_alias_items(summary_data, "court_orders", "directions")
    return orders


def _extract_exhibits(summary_data: Mapping[str, JSONValue]) -> list[JSONObject]:
    return _collect_alias_items(summary_data, "exhibits", "evidence", "documents")


def _trim_atom(text: str) -> Optional[str]:
    normalized = text.strip()
    if not normalized or len(normalized) < 8:
        return None
    sanitized = re.sub(r"\s+", " ", normalized)
    return sanitized[:280]


def _merge_usage(target: dict[str, dict[str, int]], stage: str, usage: dict[str, int]) -> None:
    stage_bucket = target.setdefault(stage, {})
    for key, value in usage.items():
        stage_bucket[key] = stage_bucket.get(key, 0) + value


def _markdown_paragraphs(markdown_text: str) -> list[str]:
    lines = markdown_text.splitlines()
    buffer: list[str] = []
    paragraphs: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                paragraphs.append(" ".join(buffer).strip())
                buffer.clear()
            continue
        if stripped.startswith("##"):
            if buffer:
                paragraphs.append(" ".join(buffer).strip())
                buffer.clear()
            paragraphs.append(stripped)
        else:
            buffer.append(stripped)
    if buffer:
        paragraphs.append(" ".join(buffer).strip())
    return paragraphs


@dataclass(slots=True)
class ComposeInputs:
    summary_markdown: str
    summary_data: JSONObject
    timeline_seeds: list[JSONObject]
    entity_hints: JSONObject
    intake: JSONObject
    case_metadata: JSONObject


@dataclass(slots=True)
class ComposeContext:
    parties: list[JSONObject] = field(default_factory=list)
    issues: list[JSONObject] = field(default_factory=list)
    facts: list[JSONObject] = field(default_factory=list)
    events: list[JSONObject] = field(default_factory=list)
    deadlines: list[JSONObject] = field(default_factory=list)
    orders: list[JSONObject] = field(default_factory=list)
    exhibits: list[JSONObject] = field(default_factory=list)
    procedural: JSONObject = field(default_factory=dict)
    claimable_atoms: list[str] = field(default_factory=list)


@dataclass(slots=True)
class GuardReport:
    ok: bool
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    checks: JSONObject = field(default_factory=dict)


@dataclass(slots=True)
class LaneAttempt:
    attempt_number: int
    document: str
    structure: GuardReport
    compliance: GuardReport
    factuality: GuardReport


@dataclass(slots=True)
class LaneOutcome:
    document: str
    structure_report: GuardReport
    compliance_report: GuardReport
    factuality_report: GuardReport
    attempts: int
    history: list[LaneAttempt]
    stage_usage: dict[str, dict[str, int]]
    token_usage: dict[str, int]
    providers: list[str]


@dataclass(slots=True)
class LaneRuntimeState:
    lane: str
    config: LaneConfig
    max_attempts: int
    attempts: int = 0
    revision_brief: Optional[str] = None
    last_document_hash: Optional[int] = None
    document: Optional[str] = None
    structure_report: Optional[GuardReport] = None
    compliance_report: Optional[GuardReport] = None
    factuality_report: Optional[GuardReport] = None
    history: list[LaneAttempt] = field(default_factory=list)
    stage_usage: dict[str, dict[str, int]] = field(default_factory=dict)
    token_usage: dict[str, int] = field(default_factory=dict)
    providers: list[str] = field(default_factory=list)

    def record_usage(self, stage: str, usage: Mapping[str, int]) -> None:
        _merge_usage(self.stage_usage, stage, dict(usage))
        for key, value in usage.items():
            self.token_usage[key] = self.token_usage.get(key, 0) + value

    def to_outcome(self) -> LaneOutcome:
        if self.document is None or self.structure_report is None or self.compliance_report is None or self.factuality_report is None:
            raise ComposeStageError(f"compose.{self.lane}", "Lane outcome requested before completion")
        return LaneOutcome(
            document=self.document,
            structure_report=_clone_report(self.structure_report),
            compliance_report=_clone_report(self.compliance_report),
            factuality_report=_clone_report(self.factuality_report),
            attempts=self.attempts,
            history=list(self.history),
            stage_usage={stage: dict(values) for stage, values in self.stage_usage.items()},
            token_usage=dict(self.token_usage),
            providers=list(self.providers),
        )


@dataclass(slots=True)
class QAReviewerResult:
    status: str
    alerts: list[str]
    recommendations: list[str]
    staff_report: str


@dataclass(slots=True)
class ComposeArtifacts:
    client_markdown: Optional[Path] = None
    lawyer_markdown: Optional[Path] = None
    client_docx: Optional[Path] = None
    lawyer_docx: Optional[Path] = None
    bundle_path: Optional[Path] = None
    qa_report: Optional[Path] = None
    staff_report: Optional[Path] = None


@dataclass(slots=True)
class ComposeResult:
    status: str
    artifacts: ComposeArtifacts
    meta_json: Path
    audit_jsonl: Path
    provider_chain: list[str]
    stage_usage: dict[str, dict[str, int]]


@dataclass(slots=True)
class ComposeConfig:
    provider_chain: list[str] = field(default_factory=lambda: list(DEFAULT_PROVIDER_CHAIN))
    temperature: float = DEFAULT_TEMPERATURE
    lawyer_temperature: float = DEFAULT_LAWYER_TEMPERATURE
    max_output_tokens: int = DEFAULT_MAX_OUTPUT_TOKENS
    max_client_attempts: int = 2
    max_lawyer_attempts: int = 2
    min_timestamp_references: int = 3
    qa_required: bool = True
    debug: bool = False
    doc_template_path: Optional[Path] = None

    @classmethod
    def from_env(cls) -> "ComposeConfig":
        providers_env = os.getenv("COMPOSE_PROVIDER_CHAIN", "")
        providers = _normalize_providers(providers_env.split(",")) if providers_env else list(DEFAULT_PROVIDER_CHAIN)
        temperature = _safe_float(os.getenv("COMPOSE_TEMPERATURE"), DEFAULT_TEMPERATURE)
        lawyer_temperature = _safe_float(os.getenv("COMPOSE_LAWYER_TEMPERATURE"), DEFAULT_LAWYER_TEMPERATURE)
        max_tokens = _safe_int(os.getenv("COMPOSE_MAX_OUTPUT_TOKENS"), DEFAULT_MAX_OUTPUT_TOKENS)
        max_client_attempts = _safe_int(os.getenv("COMPOSE_MAX_CLIENT_ATTEMPTS"), 2)
        max_lawyer_attempts = _safe_int(os.getenv("COMPOSE_MAX_LAWYER_ATTEMPTS"), 2)
        min_timestamp_references = _safe_int(os.getenv("COMPOSE_MIN_TIMESTAMP_REFERENCES"), 3)
        qa_required = _truthy(os.getenv("COMPOSE_QA_REQUIRED"), True)
        debug = _truthy(os.getenv("DEBUG"), False)
        template_env = os.getenv(DOC_TEMPLATE_ENV)
        template_path = Path(template_env).resolve() if template_env else None
        if template_path and not template_path.exists():
            logger.warning("compose.doc_template.missing", extra={"path": str(template_path)})
            template_path = None
        if not providers:
            providers = list(DEFAULT_PROVIDER_CHAIN)
        return cls(
            provider_chain=providers,
            temperature=temperature,
            lawyer_temperature=lawyer_temperature,
            max_output_tokens=max_tokens,
            max_client_attempts=max_client_attempts,
            max_lawyer_attempts=max_lawyer_attempts,
            min_timestamp_references=min_timestamp_references,
            qa_required=qa_required,
            debug=debug,
            doc_template_path=template_path,
        )


@dataclass(slots=True)
class ComposeState:
    inputs: ComposeInputs
    client: LaneRuntimeState
    lawyer: LaneRuntimeState
    context: Optional[ComposeContext] = None
    lanes: dict[str, LaneOutcome] = field(default_factory=dict)
    qa: Optional[QAReviewerResult] = None
    stage_usage: dict[str, dict[str, int]] = field(default_factory=dict)


class ComposeAgent:
    def __init__(self, config: Optional[ComposeConfig] = None) -> None:
        self.config = config or ComposeConfig.from_env()
        self.settings: LLMSettings = load_llm_settings()
        self.logger = logger

    def compose(
        self,
        *,
        case_id: str,
        case_dir: Path,
        job_id: str,
        summary_json_path: Optional[Path],
        summary_markdown_path: Optional[Path],
        transcript_path: Optional[Path],  # unused but kept for API parity
        timeline_seed_path: Optional[Path] = None,
        entity_hint_path: Optional[Path] = None,
        intake: Optional[Mapping[str, Any]] = None,
        case_metadata: Optional[Mapping[str, Any]] = None,
        provider_credentials: Optional[Mapping[str, Mapping[str, Any]]] = None,
        progress_callback: Optional[Callable[[str, str, JSONObject], None]] = None,
    ) -> ComposeResult:
        case_dir = Path(case_dir)
        docs_dir = case_dir / "docs"
        ops_dir = case_dir / "ops"
        ensure_dir(docs_dir)
        ensure_dir(ops_dir)

        summary_markdown = _read_text(summary_markdown_path)
        summary_data = _read_json(summary_json_path)
        timeline_seeds = _load_sequence(timeline_seed_path)
        entity_hints = _read_json(entity_hint_path)
        inputs = ComposeInputs(
            summary_markdown=summary_markdown,
            summary_data=summary_data,
            timeline_seeds=timeline_seeds,
            entity_hints=entity_hints,
            intake=coerce_json_object(intake) if intake else {},
            case_metadata=coerce_json_object(case_metadata) if case_metadata else {},
        )

        state = ComposeState(
            inputs=inputs,
            client=LaneRuntimeState(
                lane="client",
                config=LANE_CONFIGS["client"],
                max_attempts=self.config.max_client_attempts,
            ),
            lawyer=LaneRuntimeState(
                lane="lawyer",
                config=LANE_CONFIGS["lawyer"],
                max_attempts=self.config.max_lawyer_attempts,
            ),
        )

        provider_credentials_map: Mapping[str, JSONObject] = {
            name: coerce_json_object(payload)
            for name, payload in (provider_credentials or {}).items()
        }

        self._run_graph(
            state=state,
            provider_credentials=provider_credentials_map,
            progress=progress_callback,
        )

        if state.qa is None:
            raise ComposeStageError("compose.qa_reviewer", "QA reviewer did not execute")

        if self.config.qa_required and state.qa.status.lower() not in QA_REVIEWER_STATUS_OK:
            raise ComposeStageError(
                "compose.qa_reviewer",
                f"QA reviewer returned status '{state.qa.status}'",
            )

        artifacts = self._write_artifacts(
            state=state,
            docs_dir=docs_dir,
            job_id=job_id,
        )

        meta_payload = {
            "case_id": case_id,
            "job_id": job_id,
            "client_attempts": state.lanes["client"].attempts,
            "lawyer_attempts": state.lanes["lawyer"].attempts,
            "client_reports": _lane_history_payload(state.lanes["client"].history),
            "lawyer_reports": _lane_history_payload(state.lanes["lawyer"].history),
            "client_providers": state.lanes["client"].providers,
            "lawyer_providers": state.lanes["lawyer"].providers,
            "client_token_usage": state.lanes["client"].token_usage,
            "lawyer_token_usage": state.lanes["lawyer"].token_usage,
            "qa_status": state.qa.status,
            "qa_alerts": state.qa.alerts,
            "qa_recommendations": state.qa.recommendations,
            "provider_chain": list(self.config.provider_chain),
            "stage_usage": {stage: dict(values) for stage, values in state.stage_usage.items()},
            "bundle_path": str(artifacts.bundle_path) if artifacts.bundle_path else None,
            "client_markdown": str(artifacts.client_markdown) if artifacts.client_markdown else None,
            "lawyer_markdown": str(artifacts.lawyer_markdown) if artifacts.lawyer_markdown else None,
            "client_docx": str(artifacts.client_docx) if artifacts.client_docx else None,
            "lawyer_docx": str(artifacts.lawyer_docx) if artifacts.lawyer_docx else None,
            "qa_report": str(artifacts.qa_report) if artifacts.qa_report else None,
            "staff_report": str(artifacts.staff_report) if artifacts.staff_report else None,
            "status": "ok",
        }

        meta_json = ops_dir / f"{job_id}__compose_log.json"
        write_json_object(meta_json, meta_payload)

        audit_jsonl = ops_dir / "ops_compose.jsonl"
        append_jsonl(
            audit_jsonl,
            {
                "case_id": case_id,
                "job_id": job_id,
                "status": "ok",
                "bundle_path": str(artifacts.bundle_path) if artifacts.bundle_path else None,
                "client_markdown": str(artifacts.client_markdown) if artifacts.client_markdown else None,
                "lawyer_markdown": str(artifacts.lawyer_markdown) if artifacts.lawyer_markdown else None,
                "staff_report": str(artifacts.staff_report) if artifacts.staff_report else None,
            },
        )

        return ComposeResult(
            status="ok",
            artifacts=artifacts,
            meta_json=meta_json,
            audit_jsonl=audit_jsonl,
            provider_chain=list(self.config.provider_chain),
            stage_usage={stage: dict(values) for stage, values in state.stage_usage.items()},
        )

    # ------------------------------------------------------------------
    # Orchestration

    def _run_graph(
        self,
        *,
        state: ComposeState,
        provider_credentials: Mapping[str, JSONObject],
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> None:
        graph: Any = StateGraph(ComposeState)

        def context_node(current: ComposeState) -> ComposeState:
            return self._context_assembler(current, progress)

        def client_composer_node(current: ComposeState) -> ComposeState:
            return self._draft_lane(
                current,
                lane="client",
                provider_credentials=provider_credentials,
                progress=progress,
            )

        def client_structure_node(current: ComposeState) -> ComposeState:
            return self._structure_guard(current, lane="client", progress=progress)

        def client_compliance_node(current: ComposeState) -> ComposeState:
            return self._compliance_guard(current, lane="client", progress=progress)

        def client_factual_node(current: ComposeState) -> ComposeState:
            return self._factuality_gate(current, lane="client", progress=progress)

        def client_revision_node(current: ComposeState) -> ComposeState:
            return self._prepare_revision(current, lane="client", progress=progress)

        def lawyer_composer_node(current: ComposeState) -> ComposeState:
            return self._draft_lane(
                current,
                lane="lawyer",
                provider_credentials=provider_credentials,
                progress=progress,
            )

        def lawyer_structure_node(current: ComposeState) -> ComposeState:
            return self._structure_guard(current, lane="lawyer", progress=progress)

        def lawyer_compliance_node(current: ComposeState) -> ComposeState:
            return self._compliance_guard(current, lane="lawyer", progress=progress)

        def lawyer_factual_node(current: ComposeState) -> ComposeState:
            return self._factuality_gate(current, lane="lawyer", progress=progress)

        def lawyer_revision_node(current: ComposeState) -> ComposeState:
            return self._prepare_revision(current, lane="lawyer", progress=progress)

        def release_node(current: ComposeState) -> ComposeState:
            return self._release_gate(
                current,
                provider_credentials=provider_credentials,
                progress=progress,
            )

        graph.add_node("ContextAssembler", context_node)
        graph.add_node("ClientComposer", client_composer_node)
        graph.add_node("ClientStructureValidator", client_structure_node)
        graph.add_node("ClientComplianceGuard", client_compliance_node)
        graph.add_node("ClientFactualityGate", client_factual_node)
        graph.add_node("ClientRevision", client_revision_node)

        graph.add_node("LawyerComposer", lawyer_composer_node)
        graph.add_node("LawyerStructureValidator", lawyer_structure_node)
        graph.add_node("LawyerComplianceGuard", lawyer_compliance_node)
        graph.add_node("LawyerFactualityGate", lawyer_factual_node)
        graph.add_node("LawyerRevision", lawyer_revision_node)

        graph.add_node("ReleaseGate", release_node)

        graph.add_edge(START, "ContextAssembler")
        graph.add_edge("ContextAssembler", "ClientComposer")
        graph.add_edge("ContextAssembler", "LawyerComposer")
        graph.add_edge("ClientComposer", "ClientStructureValidator")
        graph.add_edge("ClientStructureValidator", "ClientComplianceGuard")
        graph.add_edge("ClientComplianceGuard", "ClientFactualityGate")

        def client_needs_revision(current: ComposeState) -> str:
            lane_state = current.client
            reports = [
                lane_state.structure_report,
                lane_state.compliance_report,
                lane_state.factuality_report,
            ]
            bad = [report for report in reports if report is not None and not report.ok]
            if bad and lane_state.attempts < lane_state.max_attempts:
                return "ClientRevision"
            return "LawyerComposer"

        graph.add_conditional_edges(
            "ClientFactualityGate",
            client_needs_revision,
            {
                "ClientRevision": "ClientRevision",
                "LawyerComposer": "LawyerComposer",
            },
        )
        graph.add_edge("ClientRevision", "ClientComposer")

        graph.add_edge("LawyerComposer", "LawyerStructureValidator")
        graph.add_edge("LawyerStructureValidator", "LawyerComplianceGuard")
        graph.add_edge("LawyerComplianceGuard", "LawyerFactualityGate")

        def lawyer_needs_revision(current: ComposeState) -> str:
            lane_state = current.lawyer
            reports = [
                lane_state.structure_report,
                lane_state.compliance_report,
                lane_state.factuality_report,
            ]
            bad = [report for report in reports if report is not None and not report.ok]
            if bad and lane_state.attempts < lane_state.max_attempts:
                return "LawyerRevision"
            return "ReleaseGate"

        graph.add_conditional_edges(
            "LawyerFactualityGate",
            lawyer_needs_revision,
            {
                "LawyerRevision": "LawyerRevision",
                "ReleaseGate": "ReleaseGate",
            },
        )
        graph.add_edge("LawyerRevision", "LawyerComposer")
        graph.add_edge("ReleaseGate", END)

        compiled = graph.compile()
        compiled.invoke(state)

    @staticmethod
    def _lane_state(state: ComposeState, lane: str) -> LaneRuntimeState:
        if lane == "client":
            return state.client
        if lane == "lawyer":
            return state.lawyer
        raise ComposeStageError(f"compose.{lane}", "Unknown lane")

    def _context_assembler(
        self,
        state: ComposeState,
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        stage_name = "compose.context"
        _emit(progress, stage_name, "start", {})
        state.context = _assemble_context(state.inputs, self.config.min_timestamp_references)
        _emit(progress, stage_name, "complete", {})
        return state

    def _draft_lane(
        self,
        state: ComposeState,
        *,
        lane: str,
        provider_credentials: Mapping[str, JSONObject],
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        lane_state = self._lane_state(state, lane)
        context = state.context
        if context is None:
            raise ComposeStageError(f"compose.{lane}.draft", "Compose context missing")
        is_revision = lane_state.revision_brief is not None
        stage_name = f"compose.{lane}.{'revise' if is_revision else 'draft'}"
        next_attempt = lane_state.attempts + 1
        if next_attempt > lane_state.max_attempts:
            raise ComposeStageError(stage_name, "Maximum attempts exhausted")
        _emit(progress, stage_name, "start", {"attempt": next_attempt})
        lane_state.attempts = next_attempt
        temperature = self.config.temperature if lane == "client" else self.config.lawyer_temperature
        if is_revision:
            temperature = lane_state.config.revision_temperature
        document, usage, provider = self._invoke_llm(
            stage=stage_name,
            system_prompt=lane_system_prompt(lane, revision=is_revision),
            user_prompt=_lane_user_prompt(lane, context, lane_state.revision_brief),
            temperature=temperature,
            provider_credentials=provider_credentials,
        )
        lane_state.revision_brief = None
        lane_state.document = document
        doc_hash = hash(document.strip().lower())
        if is_revision and lane_state.last_document_hash is not None and doc_hash == lane_state.last_document_hash:
            raise ComposeStageError(stage_name, "Revision produced no changes")
        lane_state.last_document_hash = doc_hash
        lane_state.structure_report = None
        lane_state.compliance_report = None
        lane_state.factuality_report = None
        lane_state.providers.append(provider)
        lane_state.record_usage(stage_name, usage)
        _merge_usage(state.stage_usage, stage_name, usage)
        _emit(
            progress,
            stage_name,
            "complete",
            {"attempt": lane_state.attempts, "provider": provider},
        )
        return state

    def _structure_guard(
        self,
        state: ComposeState,
        *,
        lane: str,
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        lane_state = self._lane_state(state, lane)
        document = lane_state.document
        if document is None:
            raise ComposeStageError(f"compose.{lane}.structure", "No draft available")
        stage_name = f"compose.{lane}.structure"
        _emit(progress, stage_name, "start", {"attempt": lane_state.attempts})
        report = _markdown_structure_report(
            document,
            lane_state.config.headings,
            min_words=lane_state.config.min_words,
        )
        if lane_state.config.readability_limit is not None:
            readability = _sentence_length_report(
                document,
                max_average_words=lane_state.config.readability_limit,
            )
            if not readability.ok:
                report.errors.extend(readability.errors)
                report.warnings.extend(readability.warnings)
        lane_state.structure_report = report
        _emit(progress, stage_name, "complete", {"ok": report.ok})
        return state

    def _compliance_guard(
        self,
        state: ComposeState,
        *,
        lane: str,
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        lane_state = self._lane_state(state, lane)
        document = lane_state.document
        if document is None:
            raise ComposeStageError(f"compose.{lane}.compliance", "No draft available")
        stage_name = f"compose.{lane}.compliance"
        _emit(progress, stage_name, "start", {"attempt": lane_state.attempts})
        report = _compliance_report(document)
        lane_state.compliance_report = report
        _emit(progress, stage_name, "complete", {"ok": report.ok})
        return state

    def _factuality_gate(
        self,
        state: ComposeState,
        *,
        lane: str,
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        lane_state = self._lane_state(state, lane)
        document = lane_state.document
        context = state.context
        if document is None or context is None:
            raise ComposeStageError(f"compose.{lane}.factuality", "Missing draft or context")
        stage_name = f"compose.{lane}.factuality"
        _emit(progress, stage_name, "start", {"attempt": lane_state.attempts})
        report = _factuality_report(
            document,
            claimable_atoms=context.claimable_atoms,
            timeline_events=context.events,
            min_timestamp_references=lane_state.config.min_timestamp_references,
        )
        lane_state.factuality_report = report
        structure_snapshot = (
            _clone_report(lane_state.structure_report)
            if lane_state.structure_report
            else GuardReport(ok=False, errors=["Structure report missing"])
        )
        compliance_snapshot = (
            _clone_report(lane_state.compliance_report)
            if lane_state.compliance_report
            else GuardReport(ok=False, errors=["Compliance report missing"])
        )
        attempt_record = LaneAttempt(
            attempt_number=lane_state.attempts,
            document=document,
            structure=structure_snapshot,
            compliance=compliance_snapshot,
            factuality=_clone_report(report),
        )
        lane_state.history.append(attempt_record)
        payload: JSONObject = {"ok": report.ok}
        if report.errors:
            payload["errors"] = list(report.errors)
        _emit(progress, stage_name, "complete", payload)
        if (
            lane_state.structure_report
            and lane_state.structure_report.ok
            and lane_state.compliance_report
            and lane_state.compliance_report.ok
            and report.ok
        ):
            state.lanes[lane] = lane_state.to_outcome()
        return state

    def _prepare_revision(
        self,
        state: ComposeState,
        *,
        lane: str,
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        lane_state = self._lane_state(state, lane)
        stage_name = f"compose.{lane}.revision"
        _emit(progress, stage_name, "start", {"attempt": lane_state.attempts})
        if lane_state.structure_report is None or lane_state.compliance_report is None or lane_state.factuality_report is None:
            raise ComposeStageError(stage_name, "Revision requested before guard reports computed")
        revision_brief = _build_revision_brief(
            lane,
            lane_state.structure_report,
            lane_state.compliance_report,
            lane_state.factuality_report,
        )
        lane_state.revision_brief = revision_brief
        _emit(
            progress,
            stage_name,
            "complete",
            {"has_revision": bool(revision_brief.strip())},
        )
        return state

    def _release_gate(
        self,
        state: ComposeState,
        *,
        provider_credentials: Mapping[str, JSONObject],
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> ComposeState:
        stage_name = "compose.release_gate"
        _emit(progress, stage_name, "start", {})
        client_outcome = state.lanes.get("client")
        lawyer_outcome = state.lanes.get("lawyer")
        if client_outcome is None or not (
            client_outcome.structure_report.ok
            and client_outcome.compliance_report.ok
            and client_outcome.factuality_report.ok
        ):
            raise ComposeStageError(stage_name, "Client lane did not pass all guards")
        if lawyer_outcome is None or not (
            lawyer_outcome.structure_report.ok
            and lawyer_outcome.compliance_report.ok
            and lawyer_outcome.factuality_report.ok
        ):
            raise ComposeStageError(stage_name, "Lawyer lane did not pass all guards")
        self._execute_qa(
            state=state,
            provider_credentials=provider_credentials,
            progress=progress,
        )
        qa_status = state.qa.status if state.qa else "missing"
        _emit(progress, stage_name, "complete", {"qa_status": qa_status})
        return state

    def _execute_qa(
        self,
        *,
        state: ComposeState,
        provider_credentials: Mapping[str, JSONObject],
        progress: Optional[Callable[[str, str, JSONObject], None]],
    ) -> None:
        stage_name = "compose.qa_reviewer"
        _emit(progress, stage_name, "start", {})

        if state.context is None:
            raise ComposeStageError(stage_name, "Compose context missing")

        payload = {
            "compose_context": state.context.procedural,
            "claimable_atoms": state.context.claimable_atoms,
            "client_brief": state.lanes["client"].document,
            "lawyer_brief": state.lanes["lawyer"].document,
        }
        user_prompt = json.dumps(payload, ensure_ascii=False)
        response, usage, provider = self._invoke_llm(
            stage=stage_name,
            system_prompt=QA_REVIEWER_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            temperature=0.0,
            provider_credentials=provider_credentials,
        )
        _merge_usage(state.stage_usage, stage_name, usage)

        try:
            parsed = json.loads(response)
        except json.JSONDecodeError as exc:  # pragma: no cover - defensive
            raise ComposeStageError(stage_name, f"Invalid QA reviewer response: {exc}") from exc

        status = coerce_str(parsed.get("status")) or "unknown"
        alerts = [coerce_str(item) or "" for item in parsed.get("alerts", []) if isinstance(item, str)]
        recommendations = [
            coerce_str(item) or ""
            for item in parsed.get("recommendations", [])
            if isinstance(item, str)
        ]
        staff_report = coerce_str(parsed.get("staff_report")) or ""
        if not staff_report.strip():
            staff_report = "# Staff Report\n\nNo staff report returned."

        state.qa = QAReviewerResult(
            status=status,
            alerts=[item for item in alerts if item],
            recommendations=[item for item in recommendations if item],
            staff_report=staff_report,
        )
        _emit(progress, stage_name, "complete", {"status": status, "provider": provider})

    # ------------------------------------------------------------------
    # Artifact writing

    def _write_artifacts(
        self,
        state: ComposeState,
        docs_dir: Path,
        job_id: str,
    ) -> ComposeArtifacts:
        artifacts = ComposeArtifacts()

        client_md_path = next_versioned(docs_dir / f"{job_id}__compose_client_v1.md")
        client_md_path.write_text(state.lanes["client"].document, encoding="utf-8")
        artifacts.client_markdown = client_md_path

        lawyer_md_path = next_versioned(docs_dir / f"{job_id}__compose_lawyer_v1.md")
        lawyer_md_path.write_text(state.lanes["lawyer"].document, encoding="utf-8")
        artifacts.lawyer_markdown = lawyer_md_path

        bundle_path = next_versioned(docs_dir / f"{job_id}__compose_bundle_v1.md")
        bundle = self._build_bundle(state)
        bundle_path.write_text(bundle, encoding="utf-8")
        artifacts.bundle_path = bundle_path

        qa_result = state.qa
        if qa_result is None:
            raise ComposeStageError("compose.write_artifacts", "QA results missing")

        staff_report_path = next_versioned(docs_dir / f"{job_id}__compose_staff_report_v1.md")
        staff_report_path.write_text(qa_result.staff_report, encoding="utf-8")
        artifacts.staff_report = staff_report_path

        qa_report_path = next_versioned(docs_dir / f"{job_id}__compose_qa_report_v1.md")
        qa_report_path.write_text(_render_qa_markdown(state), encoding="utf-8")
        artifacts.qa_report = qa_report_path

        artifacts.client_docx = self._write_docx(
            markdown=state.lanes["client"].document,
            output_prefix=docs_dir / f"{job_id}__compose_client_v1",
        )
        artifacts.lawyer_docx = self._write_docx(
            markdown=state.lanes["lawyer"].document,
            output_prefix=docs_dir / f"{job_id}__compose_lawyer_v1",
        )

        return artifacts

    def _write_docx(self, *, markdown: str, output_prefix: Path) -> Optional[Path]:
        output_path = next_versioned(output_prefix.with_suffix(".docx"))
        if self.config.doc_template_path and Document is not None and self.config.doc_template_path.exists():
            if _render_docx_from_template(self.config.doc_template_path, markdown, output_path):
                return output_path
        paragraphs = _markdown_paragraphs(markdown)
        write_basic_docx(paragraphs=paragraphs, output_path=output_path, title=output_prefix.name)
        return output_path

    def _build_bundle(self, state: ComposeState) -> str:
        sections = [
            "Part 1 – Client Summary",
            state.lanes["client"].document.strip(),
            "",
            "---",
            "",
            "Part 2 – Lawyer Brief",
            state.lanes["lawyer"].document.strip(),
        ]
        return "\n".join(sections).strip() + "\n"

    # ------------------------------------------------------------------
    # LLM Invocation

    def _invoke_llm(
        self,
        *,
        stage: str,
        system_prompt: str,
        user_prompt: str,
        temperature: float,
        provider_credentials: Mapping[str, JSONObject],
    ) -> tuple[str, dict[str, int], str]:
        providers = _normalize_providers(self.config.provider_chain)
        last_error: Optional[Exception] = None
        model_override = STAGE_MODEL_DEFAULTS.get(stage, "")

        for provider_name in providers:
            provider_meta = self.settings.provider(provider_name)
            if provider_meta is None:
                continue
            provider_info = cast(Any, provider_meta)
            default_model = cast(str, getattr(provider_info, "default_model", ""))
            model_name = model_override or default_model
            if not model_name:
                last_error = ComposeStageError(stage, f"No model configured for provider '{provider_name}'")
                continue
            credentials = provider_credentials.get(provider_name)
            try:
                runtime = build_provider_runtime_config(
                    provider=provider_meta,
                    model_name=model_name,
                    credential_payload=credentials,
                    options=None,
                )
            except ChatClientError as exc:
                last_error = exc
                continue

            try:
                client = build_chat_client(provider_runtime=runtime)
            except ChatClientError as exc:
                last_error = exc
                continue

            try:
                content, usage = client.chat(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    temperature=temperature,
                    max_tokens=self.config.max_output_tokens,
                    response_format=None,
                )
            except ChatClientError as exc:
                last_error = exc
                continue

            usage_map = {key: value for key, value in usage.items() if isinstance(value, int)}
            return content, usage_map, provider_name

        raise ComposeStageError(stage, str(last_error) if last_error else "No provider available")


# ----------------------------------------------------------------------
# Guard helpers


def _clone_report(report: GuardReport) -> GuardReport:
    return GuardReport(
        ok=report.ok,
        errors=list(report.errors),
        warnings=list(report.warnings),
        checks=dict(report.checks),
    )


def _sentence_length_report(document: str, *, max_average_words: float) -> GuardReport:
    filtered_lines = [line for line in document.splitlines() if not line.strip().startswith(('- ', '* '))]
    body = "\n".join(filtered_lines)
    body = re.sub(r"^##.+$", "", body, flags=re.MULTILINE)
    candidates = [candidate.strip() for candidate in re.split(r"(?<=[.!?])\s+|\n", body) if candidate.strip()]
    lengths: list[int] = []
    for candidate in candidates:
        words = re.findall(r"\b\w+\b", candidate)
        if words:
            lengths.append(len(words))
    if not lengths:
        return GuardReport(ok=True, errors=[], warnings=[], checks={"average_words": 0.0})
    average = sum(lengths) / float(len(lengths))
    if average > max_average_words:
        return GuardReport(
            ok=False,
            errors=[f"Average sentence length {average:.1f} exceeds {max_average_words:.0f}"],
            warnings=[],
            checks={"average_words": average},
        )
    return GuardReport(ok=True, errors=[], warnings=[], checks={"average_words": average})


_ADVICE_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"\byou should\b", re.IGNORECASE),
    re.compile(r"\byou must\b", re.IGNORECASE),
    re.compile(r"\bi recommend\b", re.IGNORECASE),
    re.compile(r"\bi advise\b", re.IGNORECASE),
    re.compile(r"\bwe recommend\b", re.IGNORECASE),
    re.compile(r"\bwe advise\b", re.IGNORECASE),
)


_RISKY_PATTERNS: tuple[tuple[re.Pattern[str], str], ...] = (
    (re.compile(r"\bshould consider\b", re.IGNORECASE), "Suggestive wording"),
    (re.compile(r"\blikely\b", re.IGNORECASE), "Speculative wording"),
    (re.compile(r"\bprobably\b", re.IGNORECASE), "Speculative wording"),
)


_HEADING_PATTERN = re.compile(r"^##\s+(.+?)\s*$", re.MULTILINE)


def _split_markdown_sections(document: str) -> list[tuple[str, str]]:
    normalized = document.replace("\r\n", "\n")
    matches = list(_HEADING_PATTERN.finditer(normalized))
    sections: list[tuple[str, str]] = []
    for index, match in enumerate(matches):
        heading = "## " + match.group(1).strip()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(normalized)
        sections.append((heading, normalized[start:end].strip()))
    return sections


def _markdown_structure_report(document: str, required: Sequence[str], *, min_words: int) -> GuardReport:
    text = document.strip()
    if not text:
        return GuardReport(ok=False, errors=["Document is empty"], warnings=[], checks={})

    sections = _split_markdown_sections(text)
    heading_order = [heading for heading, _ in sections]
    errors: list[str] = []
    warnings: list[str] = []

    last_index = -1
    for heading in required:
        occurrences = [idx for idx, present in enumerate(heading_order) if present == heading]
        if not occurrences:
            errors.append(f"Missing heading '{heading}'")
            continue
        if len(occurrences) > 1:
            errors.append(f"Duplicate heading '{heading}'")
        current_index = occurrences[0]
        if current_index <= last_index:
            errors.append(f"Heading '{heading}' out of order")
        last_index = current_index

    allowed = set(required)
    for heading in heading_order:
        if heading not in allowed:
            warnings.append(f"Unexpected heading '{heading}'")

    section_map = {heading: content for heading, content in sections}
    for heading in required:
        content = section_map.get(heading, "")
        word_count = len(re.findall(r"\b\w+\b", content))
        if word_count < min_words:
            errors.append(
                f"Section '{heading}' too short (has {word_count} words, expected at least {min_words})"
            )

    return GuardReport(ok=not errors, errors=errors, warnings=warnings, checks={})


def _compliance_report(document: str) -> GuardReport:
    errors: list[str] = []
    warnings: list[str] = []
    for line in document.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(">"):  # quoted transcript or order
            continue
        lowered = stripped.lower()
        if lowered.startswith(("order:", "court order:", "epo:", "directive:")):
            continue
        for pattern in _ADVICE_PATTERNS:
            if pattern.search(lowered):
                errors.append(f"Disallowed advice language: '{stripped}'")
        for pattern, label in _RISKY_PATTERNS:
            if pattern.search(lowered):
                warnings.append(f"{label}: '{stripped}'")
    return GuardReport(ok=not errors, errors=errors, warnings=warnings, checks={})


def _factuality_report(
    document: str,
    *,
    claimable_atoms: Sequence[str],
    timeline_events: Sequence[JSONObject],
    min_timestamp_references: int,
) -> GuardReport:
    sentences = [candidate.strip() for candidate in re.split(r"(?<=[.!?])\s+|\n", document) if candidate.strip()]
    atoms = [atom.lower() for atom in claimable_atoms if atom]
    errors: list[str] = []
    warnings: list[str] = []

    references_found = len(re.findall(r"\[(\d{2}:\d{2})\]", document))
    required_refs = min(len(list(timeline_events)), min_timestamp_references)
    if references_found < required_refs:
        errors.append(f"Found {references_found} timestamp references; expected at least {required_refs}")

    for sentence in sentences:
        lowered = sentence.lower()
        if lowered.startswith("## ") or lowered.startswith("information not provided"):
            continue
        if lowered.startswith(('- ', '* ')):
            continue
        if re.search(r"\[(\d{2}:\d{2})\]", sentence):
            continue
        if len(lowered) < 24:
            continue
        if atoms and any(atom in lowered or SequenceMatcher(None, atom, lowered).ratio() >= 0.82 for atom in atoms):
            continue
        errors.append(f"Unsupported assertion: '{sentence}'")

    event_ids = [coerce_str(event.get("id")) or "" for event in timeline_events]
    referenced_ids = {
        event_id
        for event_id in event_ids
        if event_id and re.search(rf"\b{re.escape(event_id)}\b", document)
    }
    missing_ids = [event_id for event_id in event_ids if event_id and event_id not in referenced_ids]
    if missing_ids:
        warnings.append(f"Timeline events not referenced: {', '.join(sorted(missing_ids))}")

    return GuardReport(ok=not errors, errors=errors, warnings=warnings, checks={})


def _build_revision_brief(
    lane: str,
    structure: GuardReport,
    compliance: GuardReport,
    factuality: GuardReport,
) -> str:
    sections: list[str] = []
    if structure.errors or structure.warnings:
        sections.append("Structure:")
        for item in structure.errors:
            sections.append(f"- {item}")
        for item in structure.warnings:
            sections.append(f"- (warning) {item}")
    if compliance.errors or compliance.warnings:
        sections.append("Compliance:")
        for item in compliance.errors:
            sections.append(f"- {item}")
        for item in compliance.warnings:
            sections.append(f"- (warning) {item}")
    if factuality.errors or factuality.warnings:
        sections.append("Factuality:")
        for item in factuality.errors:
            sections.append(f"- {item}")
        for item in factuality.warnings:
            sections.append(f"- (warning) {item}")
    if not sections:
        sections.append("No issues detected; maintain required style and references.")
    header = REVISION_HEADER_TEMPLATE.format(lane=lane)
    return "\n".join([header, *sections])


# ----------------------------------------------------------------------
# Context assembly


def _assemble_context(inputs: ComposeInputs, min_timestamp_references: int) -> ComposeContext:
    summary_data = inputs.summary_data
    parties = _extract_parties(summary_data)
    issues = _collect_alias_items(summary_data, "issues", "key_issues")
    facts = _collect_alias_items(summary_data, "facts", "key_facts")
    deadlines = _extract_deadlines(summary_data)
    orders = _extract_orders(summary_data)
    exhibits = _extract_exhibits(summary_data)

    events: list[JSONObject] = []
    for index, item in enumerate(inputs.timeline_seeds):
        event = coerce_json_object(item)
        event_id = coerce_str(event.get("id")) or f"event-{index + 1}"
        summary = coerce_str(event.get("summary")) or coerce_str(event.get("title")) or ""
        ts_start = coerce_float(event.get("ts_start"))
        ts_end = coerce_float(event.get("ts_end"))
        speakers = coerce_str_list(event.get("speakers"))
        references = coerce_str_list(event.get("references"))
        speakers_json = cast(list[JSONValue], list(speakers))
        references_json = cast(list[JSONValue], list(references))
        events.append(
            {
                "id": event_id,
                "summary": summary,
                "ts_start": ts_start,
                "ts_end": ts_end,
                "speakers": speakers_json,
                "references": references_json,
            }
        )

    for hint in coerce_object_list(inputs.entity_hints.get("entities")):
        parties.append(coerce_json_object(hint))

    claimable: set[str] = set()
    for fact in facts:
        text = coerce_str(fact.get("text")) or coerce_str(fact.get("summary")) or ""
        trimmed = _trim_atom(text)
        if trimmed:
            claimable.add(trimmed)
    for event in events:
        trimmed = _trim_atom(coerce_str(event.get("summary")) or "")
        if trimmed:
            claimable.add(trimmed)
    for deadline in deadlines:
        trimmed = _trim_atom(coerce_str(deadline.get("text")) or coerce_str(deadline.get("label")) or "")
        if trimmed:
            claimable.add(trimmed)
    for order in orders:
        trimmed = _trim_atom(coerce_str(order.get("text")) or "")
        if trimmed:
            claimable.add(trimmed)

    for line in inputs.summary_markdown.splitlines():
        if line.strip().startswith("#"):
            continue
        trimmed = _trim_atom(line)
        if trimmed:
            claimable.add(trimmed)

    procedural = _coerce_optional_object(summary_data.get("procedural"))
    if inputs.intake:
        procedural = {**procedural, **inputs.intake}
    if inputs.case_metadata:
        procedural = {**procedural, **inputs.case_metadata}

    return ComposeContext(
        parties=parties,
        issues=issues,
        facts=facts,
        events=events,
        deadlines=deadlines,
        orders=orders,
        exhibits=exhibits,
        procedural=procedural,
        claimable_atoms=sorted(claimable),
    )


def _serialize_context(context: ComposeContext) -> str:
    payload = {
        "parties": context.parties,
        "issues": context.issues,
        "facts": context.facts,
        "events": context.events,
        "deadlines": context.deadlines,
        "orders": context.orders,
        "exhibits": context.exhibits,
        "procedural": context.procedural,
        "claimable_atoms": context.claimable_atoms,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


# ----------------------------------------------------------------------
# QA Report rendering


def _render_qa_markdown(state: ComposeState) -> str:
    qa = state.qa
    assert qa is not None
    lines = ["# QA Review", ""]
    lines.append(f"**Status:** {qa.status}")
    if qa.alerts:
        lines.append("## Alerts")
        lines.extend([f"- {alert}" for alert in qa.alerts])
    if qa.recommendations:
        lines.append("")
        lines.append("## Recommendations")
        lines.extend([f"- {rec}" for rec in qa.recommendations])

    def _lane_section(lane_key: str, title: str, outcome: LaneOutcome) -> None:
        lines.append("")
        lines.append(f"## {title} Lane")
        lines.append(f"- Attempts: {outcome.attempts}")
        lines.append(f"- Structure: {'ok' if outcome.structure_report.ok else 'fail'}")
        if outcome.structure_report.errors:
            lines.extend([f"  - {error}" for error in outcome.structure_report.errors])
        lines.append(f"- Compliance: {'ok' if outcome.compliance_report.ok else 'fail'}")
        if outcome.compliance_report.errors:
            lines.extend([f"  - {error}" for error in outcome.compliance_report.errors])
        lines.append(f"- Factuality: {'ok' if outcome.factuality_report.ok else 'fail'}")
        if outcome.factuality_report.errors:
            lines.extend([f"  - {error}" for error in outcome.factuality_report.errors])

    _lane_section("client", "Client", state.lanes["client"])
    _lane_section("lawyer", "Lawyer", state.lanes["lawyer"])

    lines.append("")
    lines.append("## Staff Report")
    lines.append(qa.staff_report.strip())
    return "\n".join(lines).strip() + "\n"


# ----------------------------------------------------------------------
# DOCX template rendering (optional)


def _render_docx_from_template(template_path: Path, markdown: str, output_path: Path) -> bool:
    try:
        doc_obj = cast(Any, Document(str(template_path)))
    except Exception:  # pragma: no cover - template issues
        return False
    placeholder = "{{CONTENT}}"
    replaced = False
    for paragraph in doc_obj.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, markdown)
            replaced = True
    if not replaced:
        # Fallback to simple append at the end
        doc_obj.add_paragraph(markdown)
    try:
        doc_obj.save(str(output_path))
    except Exception:  # pragma: no cover - write issues
        return False
    return True


# ----------------------------------------------------------------------
# Utility


def _emit(
    progress: Optional[Callable[[str, str, JSONObject], None]],
    stage: str,
    event: str,
    payload: JSONObject,
) -> None:
    if progress is None:
        logger.debug("compose.stage", extra={"stage": stage, "event": event, "payload": payload})
        return
    try:
        progress(stage, event, payload)
    except Exception:  # pragma: no cover - defensive
        logger.debug("compose.progress_callback_failed", exc_info=True)


def _lane_history_payload(history: Sequence[LaneAttempt]) -> list[JSONObject]:
    serialized: list[JSONObject] = []
    for attempt in history:
        structure_payload: JSONObject = {
            "ok": attempt.structure.ok,
            "errors": cast(list[JSONValue], list(attempt.structure.errors)),
            "warnings": cast(list[JSONValue], list(attempt.structure.warnings)),
        }
        compliance_payload: JSONObject = {
            "ok": attempt.compliance.ok,
            "errors": cast(list[JSONValue], list(attempt.compliance.errors)),
            "warnings": cast(list[JSONValue], list(attempt.compliance.warnings)),
        }
        factual_payload: JSONObject = {
            "ok": attempt.factuality.ok,
            "errors": cast(list[JSONValue], list(attempt.factuality.errors)),
            "warnings": cast(list[JSONValue], list(attempt.factuality.warnings)),
        }
        serialized.append(
            {
                "attempt": attempt.attempt_number,
                "structure": structure_payload,
                "compliance": compliance_payload,
                "factuality": factual_payload,
            }
        )
    return serialized


def _load_sequence(path: Optional[Path]) -> list[JSONObject]:
    if path is None or not path.exists():
        return []
    try:
        payload = load_json_value(path, context=str(path))
    except ValueError:
        return []
    if isinstance(payload, Mapping):
        events = coerce_object_list(payload.get("events"))
        return [coerce_json_object(item) for item in events]
    if isinstance(payload, Sequence):
        return [coerce_json_object(item) for item in payload if isinstance(item, Mapping)]
    return []


def _read_text(path: Optional[Path]) -> str:
    if path is None:
        return ""
    try:
        return path.read_text(encoding="utf-8")
    except Exception:  # pragma: no cover - defensive
        return ""


def _read_json(path: Optional[Path]) -> JSONObject:
    if path is None or not path.exists():
        return {}
    try:
        return load_json_object(path, context=str(path))
    except Exception:  # pragma: no cover - defensive
        return {}


def _lane_user_prompt(lane: str, context: ComposeContext, revision_brief: Optional[str]) -> str:
    serialized_context = _serialize_context(context)
    payload: dict[str, JSONValue] = {
        "context": json.loads(serialized_context),
        "lane": lane,
    }
    if revision_brief:
        instruction = (
            CLIENT_REVISION_USER_INSTRUCTION if lane == "client" else LAWYER_REVISION_USER_INSTRUCTION
        )
        payload["revision_brief"] = revision_brief
        payload["instruction"] = instruction
        return json.dumps(payload, ensure_ascii=False)
    instruction = (
        CLIENT_DRAFT_USER_INSTRUCTION if lane == "client" else LAWYER_DRAFT_USER_INSTRUCTION
    )
    payload["instruction"] = instruction
    return json.dumps(payload, ensure_ascii=False)


__all__ = [
    "ComposeAgent",
    "ComposeConfig",
    "ComposeResult",
    "ComposeArtifacts",
    "ComposeStageError",
]
